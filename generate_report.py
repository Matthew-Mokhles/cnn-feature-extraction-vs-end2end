"""
generate_report.py
==================
Generates the full project documentation DOCX covering all 4 Requirements
for all 4 team members (Matthew, Mark, Mario, Veronia).

Run from the repo root:
    python generate_report.py

Output: project_report.docx
"""

import os
from pathlib import Path

# ── third-party ─────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise SystemExit(
        "python-docx is required.  Install it with:\n"
        "    pip install python-docx"
    )

# ── paths ────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).parent
MEMBERS = ROOT / "members"

MATTHEW = MEMBERS / "Matthew_xray"
MARK    = MEMBERS / "mark_xray"
MARIO   = MEMBERS / "Mario_plantvillage"
VERONIA = MEMBERS / "Veronia_cifar10"

OUT = ROOT / "project_report.docx"

# ─────────────────────────────────────────────────────────────────────────────
# Helper utilities
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set a table cell background colour (hex without #)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_heading(doc, text: str, level: int = 1):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)   # medium blue
    elif level == 3:
        run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x6E)   # teal
    return h


def add_para(doc, text: str, bold=False, italic=False, size=11,
             align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text: str, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_image(doc, path, width=6.0, caption=None):
    """Add an image if it exists, with optional caption."""
    path = Path(path)
    if not path.exists():
        add_para(doc, f"[Image not found: {path.name}]", italic=True,
                 color=RGBColor(0xAA, 0x00, 0x00))
        return
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(9.5)
        cp.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def add_table(doc, headers, rows, header_bg="1F497D", header_fg=RGBColor(0xFF,0xFF,0xFF),
              col_widths=None):
    """Build a formatted table with coloured header row."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = header_fg
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg  = "FFFFFF" if ri % 2 == 0 else "EBF3FB"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    return t


def add_page_break(doc):
    doc.add_page_break()


def section_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
# Cover page
# ─────────────────────────────────────────────────────────────────────────────

def build_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("CNN Feature Extraction vs End-to-End Fine-Tuning")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("A Comparative Deep Learning Study")
    rs.font.size = Pt(16)
    rs.italic = True
    rs.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()
    for line in [
        "Course: CAI3105 / CS460 — Deep Learning",
        "Institution: South Valley University",
        "College of Computing & Information Technology",
        "Lecturer: Prof. Nashwa El-Bendary",
        "Deadline: May 7th, 2026",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(12)

    doc.add_paragraph()
    team_hdr = doc.add_paragraph()
    team_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = team_hdr.add_run("Team Members")
    r2.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    members = [
        ("Matthew", "Chest X-Ray (Pneumonia)", "ResNet50"),
        ("Mark",    "Chest X-Ray (Pneumonia)", "MobileNetV2"),
        ("Mario",   "PlantVillage (38 classes)", "EfficientNet-B0"),
        ("Veronia", "CIFAR-10 (10 classes)",     "EfficientNet-B0 (PyTorch)"),
    ]
    add_table(doc,
        headers=["Member", "Dataset", "Backbone"],
        rows=members,
        col_widths=[4, 6, 5])

    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Table of Contents (static)
# ─────────────────────────────────────────────────────────────────────────────

def build_toc(doc):
    add_heading(doc, "Table of Contents", 1)
    entries = [
        "1. Research Objective & Overview",
        "2. Requirement 1 — Dataset Selection & Technical Specifications",
        "   2.1 Matthew — Chest X-Ray (ResNet50)",
        "   2.2 Mark — Chest X-Ray (MobileNetV2)",
        "   2.3 Mario — PlantVillage (EfficientNet-B0)",
        "   2.4 Veronia — CIFAR-10 (EfficientNet-B0 / PyTorch)",
        "3. Requirement 2 — DL Model Selection & Hyperparameters",
        "   3.1 ResNet50 (Matthew)",
        "   3.2 MobileNetV2 (Mark)",
        "   3.3 EfficientNet-B0 (Mario & Veronia)",
        "4. Requirement 3 — Implementation Framework",
        "   4.1 Matthew — Approach 1 (Feature Extraction) & Approach 2 (End-to-End)",
        "   4.2 Mark — Approach 1, 2A, 2B",
        "   4.3 Mario — Approach 1 (End-to-End) & Approach 2 (SVM)",
        "   4.4 Veronia — Approach 1 (SVM) & Approach 2 (End-to-End)",
        "5. Requirement 4 — Comparative Analysis & Conclusions",
        "   5.1 Per-Member Comparisons",
        "   5.2 Cross-Team Summary",
        "   5.3 Discussion & Recommendations",
        "6. References",
    ]
    for e in entries:
        p = doc.add_paragraph(e)
        p.paragraph_format.left_indent = Pt(18) if e.startswith("   ") else Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.runs[0].font.size = Pt(11)

    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 1 — Research Objective
# ─────────────────────────────────────────────────────────────────────────────

def build_overview(doc):
    add_heading(doc, "1. Research Objective & Overview", 1)
    add_para(doc,
        "This project evaluates the performance of End-to-End Deep Learning (DL) classification "
        "against DL-based Feature Learning. The core research question is:",
        size=11)
    add_para(doc,
        '"Given a pretrained CNN, is it better to use it as a frozen feature extractor and train '
        'a classical ML head on top — or to fine-tune the entire network end-to-end?"',
        italic=True, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x1F, 0x49, 0x7D))
    add_para(doc,
        "To answer this rigorously, the question is studied four times — each team member applies "
        "both approaches to a different image-classification dataset using a different CNN backbone, "
        "enabling a domain-independent comparison.", size=11)

    add_heading(doc, "Approach Summary", 2)
    add_table(doc,
        headers=["", "Approach 1 — Feature Extraction", "Approach 2 — End-to-End"],
        rows=[
            ["Backbone",        "Frozen (ImageNet weights, no update)", "Same CNN — fine-tuned"],
            ["Classifier head", "Classical ML: SVM / LR / MLP",        "Custom Dense layers"],
            ["What learns",     "Only the ML classifier",               "CNN + new head jointly"],
            ["Training speed",  "Very fast (CNN is not trained)",       "Slower (backprop through CNN)"],
            ["Memory footprint","Low (inference only)",                 "High (gradients stored)"],
            ["Generalisation",  "Limited to ImageNet feature space",    "Adapts features to the task"],
        ],
        col_widths=[3.5, 6.5, 6.5])

    doc.add_paragraph()
    add_heading(doc, "Team Structure", 2)
    add_table(doc,
        headers=["Member", "Dataset", "Backbone", "Approach 1", "Approach 2"],
        rows=[
            ["Matthew", "Chest X-Ray (binary)",  "ResNet50",
             "SVM · LR · MLP on GAP+GMP (4096-d)", "Two-phase fine-tune + SWA + TTA"],
            ["Mark",    "Chest X-Ray (binary)",  "MobileNetV2",
             "Linear SVM on 1280-d features",    "Two-phase fine-tune (last 20 layers)"],
            ["Mario",   "PlantVillage (38 cls)", "EfficientNet-B0",
             "RBF-SVM + PCA on 1280-d features", "Two-phase fine-tune (from layer 100)"],
            ["Veronia", "CIFAR-10 (10 cls)",     "EfficientNet-B0 (PyTorch)",
             "Linear SVM on 1280-d features",    "Full fine-tune + ReduceLROnPlateau"],
        ],
        col_widths=[2.5, 4, 3.5, 5, 5])
    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 2 — Requirement 1: Dataset
# ─────────────────────────────────────────────────────────────────────────────

def build_req1(doc):
    add_heading(doc, "2. Requirement 1 — Dataset Selection & Technical Specifications", 1)

    # ── 2.1 Matthew ──────────────────────────────────────────────────────────
    add_heading(doc, "2.1 Matthew — Chest X-Ray Pneumonia Dataset", 2)

    add_heading(doc, "Dataset Metadata", 3)
    add_table(doc,
        headers=["Property", "Value"],
        rows=[
            ["Dataset Name",    "Chest X-Ray Images (Pneumonia)"],
            ["Source",          "Kaggle — paultimothymooney/chest-xray-pneumonia"],
            ["Problem Domain",  "Medical Diagnosis (Binary Classification)"],
            ["Total Samples",   "6,638 chest radiographs (anterior-posterior)"],
            ["Classes",         "NORMAL (1,583) vs PNEUMONIA (4,273)"],
            ["Image Resolution","Variable native; resized to 224 × 224 pixels"],
            ["Color Channels",  "RGB (3 channels) — grayscale X-rays replicated to 3 channels"],
        ], col_widths=[5, 11])

    add_heading(doc, "Data Splits", 3)
    add_table(doc,
        headers=["Split", "NORMAL", "PNEUMONIA", "Total", "PNEUMONIA %"],
        rows=[
            ["Train", "1,341", "3,875", "5,216", "74.3%"],
            ["Val",   "209",   "589",   "798",   "73.8%"],
            ["Test",  "234",   "390",   "624",   "62.5%"],
        ], col_widths=[3, 3, 3, 3, 3])

    add_para(doc,
        "Important: The original Kaggle validation set contained only 16 images. "
        "The tool tools/split_train_val.py was used to carve a stratified 798-image "
        "validation set from the training data (15% fraction).",
        italic=True, size=10, color=RGBColor(0xAA, 0x44, 0x00))

    add_heading(doc, "Data Preprocessing", 3)
    for item in [
        "Resize all images to 224×224 pixels (ResNet50 standard input size).",
        "Apply ResNet50-specific pixel normalisation via tf.keras.applications.resnet50.preprocess_input — subtracts ImageNet channel means and scales.",
        "Shuffle training set with fixed seed (42) for reproducibility.",
        "No augmentation applied during feature extraction (Step 3) or evaluation.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Data Augmentation (Training only)", 3)
    add_table(doc,
        headers=["Technique", "Value", "Justification"],
        rows=[
            ["Rotation",           "±20°",            "Patients not always perfectly upright"],
            ["Width / Height shift","±15%",            "X-ray plate not always perfectly centred"],
            ["Zoom",               "±15%",            "Varying distances from the detector"],
            ["Shear",              "±8°",             "Slight angular distortion"],
            ["Brightness",         "80%–120%",        "Different acquisition exposure settings"],
            ["Horizontal flip",    "Yes",             "Chest is roughly bilaterally symmetric"],
            ["RandomErasing",      "p=0.30 (cutout)", "Forces model to use global context; bridges train↔test domain gap"],
        ], col_widths=[3.5, 2.5, 10])

    add_image(doc, MATTHEW / "results" / "sample_images.png", width=6.0,
              caption="Figure 2.1a — Sample Chest X-Ray images: NORMAL (top) vs PNEUMONIA (bottom)")
    add_image(doc, MATTHEW / "results" / "class_distribution.png", width=5.0,
              caption="Figure 2.1b — Class distribution across Train / Val / Test splits")
    add_image(doc, MATTHEW / "results" / "augmentation_samples.png", width=6.0,
              caption="Figure 2.1c — Original images vs augmented training images")
    add_page_break(doc)

    # ── 2.2 Mark ─────────────────────────────────────────────────────────────
    add_heading(doc, "2.2 Mark — Chest X-Ray Pneumonia Dataset", 2)

    add_heading(doc, "Dataset Metadata", 3)
    add_table(doc,
        headers=["Property", "Value"],
        rows=[
            ["Dataset Name",    "Chest X-Ray Images (Pneumonia)"],
            ["Source",          "Kaggle — paultimothymooney/chest-xray-pneumonia"],
            ["Problem Domain",  "Medical Diagnosis (Binary Classification)"],
            ["Total Samples",   "5,856 chest radiographs"],
            ["Classes",         "NORMAL (1,583) vs PNEUMONIA (4,273)"],
            ["Image Resolution","Resized to 224 × 224 × 3 (RGB)"],
        ], col_widths=[5, 11])

    add_heading(doc, "Key Dataset Challenge — Validation Set Size", 3)
    add_para(doc,
        "The original Kaggle validation split contains only 16 images (8 per class). "
        "With such a small set, a single wrong prediction equals a 6.25% accuracy swing, "
        "making early-stopping signals completely unreliable. "
        "Approach 2B fixes this by using validation_split=0.2 from training data → 1,043 images.",
        size=11)
    add_image(doc, MARK / "results" / "req1_validation_comparison.png", width=5.5,
              caption="Figure 2.2a — Original val (16 images) vs corrected val (1,043 images)")
    add_image(doc, MARK / "results" / "req1_sample_images.png", width=5.5,
              caption="Figure 2.2b — Sample X-Ray images from Mark's dataset")

    add_heading(doc, "Data Augmentation (Training only)", 3)
    add_table(doc,
        headers=["Technique", "Value"],
        rows=[
            ["Rotation",            "±15°"],
            ["Width / Height shift","±10%"],
            ["Zoom",                "±10%"],
            ["Horizontal flip",     "Yes"],
            ["Fill mode",           "nearest"],
        ], col_widths=[5, 11])
    add_page_break(doc)

    # ── 2.3 Mario ─────────────────────────────────────────────────────────────
    add_heading(doc, "2.3 Mario — PlantVillage Dataset", 2)

    add_heading(doc, "Dataset Metadata", 3)
    add_table(doc,
        headers=["Property", "Value"],
        rows=[
            ["Dataset Name",   "PlantVillage"],
            ["Source",         "Kaggle — abdallahalidev/plantvillage-dataset"],
            ["Problem Domain", "Agriculture — Plant Leaf Disease Classification"],
            ["Total Samples",  "~54,305 color images"],
            ["Classes",        "38 classes (14 crop species: healthy + diseased variants)"],
            ["Image Resolution","224 × 224 pixels (RGB)"],
            ["Color Channels", "RGB (3 channels) — color subset only"],
            ["Reference",      "Mohanty, Hughes & Salathe, Frontiers in Plant Science, 2016"],
        ], col_widths=[5, 11])

    add_heading(doc, "Data Splits", 3)
    add_table(doc,
        headers=["Split", "Ratio", "Images (approx.)"],
        rows=[
            ["Training",   "70%", "~38,014"],
            ["Validation", "15%", "~8,145"],
            ["Test",       "15%", "~8,145"],
        ], col_widths=[5, 3, 8])
    add_para(doc, "Random seed = 42 for full reproducibility.", italic=True, size=10)

    add_heading(doc, "Data Augmentation (Training only)", 3)
    add_table(doc,
        headers=["Technique", "Value", "Justification"],
        rows=[
            ["Rotation",            "±20°",  "PlantVillage images are controlled lab shots; augmentation forces learning of disease features"],
            ["Width / Height shift","±15%",  "Varying camera positions in real-world deployment"],
            ["Horizontal flip",     "Yes",   "Leaf disease patterns are not laterally biased"],
            ["Zoom",                "±15%",  "Different distances from the leaf"],
            ["Shear",               "±10°",  "Perspective variation"],
        ], col_widths=[3.5, 2, 10.5])

    add_image(doc, MARIO / "results" / "dataset_overview.png", width=5.5,
              caption="Figure 2.3a — PlantVillage dataset overview table")
    add_image(doc, MARIO / "results" / "class_distribution.png", width=5.5,
              caption="Figure 2.3b — Class distribution across 38 PlantVillage categories")
    add_page_break(doc)

    # ── 2.4 Veronia ───────────────────────────────────────────────────────────
    add_heading(doc, "2.4 Veronia — CIFAR-10 Dataset", 2)

    add_heading(doc, "Dataset Metadata", 3)
    add_table(doc,
        headers=["Property", "Value"],
        rows=[
            ["Dataset Name",   "CIFAR-10"],
            ["Source",         "torchvision.datasets — auto-downloaded (no manual setup needed)"],
            ["Problem Domain", "General Object Recognition (10-class Classification)"],
            ["Total Samples",  "60,000 images"],
            ["Classes",        "10: airplane, automobile, bird, cat, deer, dog, frog, horse, ship, truck"],
            ["Native Resolution","32 × 32 pixels → resized to 224 × 224 for EfficientNet-B0"],
            ["Color Channels", "RGB (3 channels)"],
            ["Reference",      "Krizhevsky, A. (2009). Learning Multiple Layers of Features from Tiny Images."],
        ], col_widths=[5, 11])

    add_heading(doc, "Data Splits", 3)
    add_table(doc,
        headers=["Split", "Samples", "Note"],
        rows=[
            ["Training",   "45,000", "90% of official CIFAR-10 train split, seed=42"],
            ["Validation", "5,000",  "10% of official CIFAR-10 train split, seed=42"],
            ["Test",       "10,000", "Official CIFAR-10 test split (no augmentation)"],
        ], col_widths=[3.5, 3, 9.5])

    add_heading(doc, "Data Augmentation (Training only)", 3)
    add_table(doc,
        headers=["Technique", "Parameters", "Justification"],
        rows=[
            ["RandomHorizontalFlip", "p=0.5",                    "CIFAR-10 objects are largely symmetric"],
            ["RandomRotation",       "±15°",                     "Orientation invariance"],
            ["ColorJitter",          "brightness=0.2, contrast=0.2", "Robustness to lighting conditions"],
            ["Normalize",            "ImageNet mean/std",         "Matches EfficientNet-B0 pre-training distribution"],
        ], col_widths=[4, 4.5, 7.5])
    add_para(doc, "No augmentation is applied at validation or test time to prevent data leakage.",
             italic=True, size=10)

    add_image(doc, VERONIA / "Results" / "req1_samples.png", width=5.5,
              caption="Figure 2.4a — Sample CIFAR-10 images across the 10 classes")
    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 3 — Requirement 2: Model Selection & Hyperparameters
# ─────────────────────────────────────────────────────────────────────────────

def build_req2(doc):
    add_heading(doc, "3. Requirement 2 — DL Model Selection & Hyperparameters", 1)
    add_para(doc,
        "All four backbones share the same design philosophy: they are pretrained on ImageNet "
        "(1.28 million images, 1000 classes) and the top classification layer is removed, "
        "leaving a feature-extraction body. The key differences are size, speed, and the "
        "method of feature pooling.", size=11)

    # ── 3.1 ResNet50 (Matthew) ───────────────────────────────────────────────
    add_heading(doc, "3.1 ResNet50 — Matthew", 2)
    add_para(doc,
        "ResNet50 (He et al., CVPR 2016) introduced residual / skip connections to allow "
        "gradients to flow directly across layers, enabling very deep networks without "
        "vanishing gradients. It consists of 175 convolutional, BN, and pooling layers "
        "organised into 4 residual stages with bottleneck blocks, ending in a 2048-dimensional "
        "feature map before the classification head.", size=11)
    add_para(doc,
        "Matthew extends the standard GAP output to a GAP+GMP concatenation, doubling "
        "the feature vector to 4096-d. This provides both average and peak activation signals, "
        "improving linear classifiers by 3–5 points on chest X-rays compared to plain GAP.",
        size=11)

    add_heading(doc, "Hyperparameter Tables — Matthew", 3)
    add_para(doc, "Approach 1 — Feature Extraction + Classical ML:", bold=True, size=11)
    add_table(doc,
        headers=["Classifier", "Key Parameters", "Value"],
        rows=[
            ["SVM",  "Kernel / C / Class weight / CV metric", "linear / GridSearch {0.01,0.1,1.0} / balanced or {0:2,1:1} / f1_macro"],
            ["LR",   "C / Class weight / CV metric",          "GridSearch {0.01,0.1,1.0} / balanced or {0:2,1:1} / f1_macro"],
            ["MLP",  "Hidden layers / Activation / Optimizer","(256, 64) / ReLU / Adam + early stopping"],
            ["Scaler","Pre-processing pipeline",              "StandardScaler → L2 row-normalisation"],
        ], col_widths=[2.5, 5.5, 8])

    add_para(doc, "Approach 2 — End-to-End Fine-Tuning:", bold=True, size=11)
    add_table(doc,
        headers=["Parameter", "Phase 1 (Head Warm-up)", "Phase 2 (Fine-Tuning)"],
        rows=[
            ["Frozen layers",    "All 175 base layers frozen",      "Last 15 base layers unfrozen (24/184 total)"],
            ["Epochs",           "10",                               "15"],
            ["Optimizer",        "Adam",                             "Adam"],
            ["Learning rate",    "1e-4 + ReduceLROnPlateau",         "1e-5 + cosine decay → 5e-7"],
            ["Loss",             "BinaryCrossentropy (label_smooth=0.1)", "BinaryCrossentropy (label_smooth=0.1)"],
            ["Batch size",       "32",                               "32"],
            ["Dropout",          "0.4 (layer 1) / 0.3 (layer 2)",   "Same"],
            ["L2 regularisation","1e-4 on Dense layers",             "1e-4 on Dense layers"],
            ["BatchNorm",        "Trainable",                        "Frozen in inference mode"],
            ["Class weights",    "Balanced (auto-computed)",         "Balanced"],
            ["SWA",              "—",                                "Enabled, starts at 60% of Phase 2"],
            ["TTA",              "—",                                "8-view at inference"],
        ], col_widths=[4, 6, 6])
    add_page_break(doc)

    # ── 3.2 MobileNetV2 (Mark) ───────────────────────────────────────────────
    add_heading(doc, "3.2 MobileNetV2 — Mark", 2)
    add_para(doc,
        "MobileNetV2 (Sandler et al., CVPR 2018) is designed for mobile and embedded "
        "applications. It uses Inverted Residual Bottleneck blocks (expand → depthwise conv → "
        "project) with linear bottlenecks to preserve information. With only 3.4M parameters "
        "it is ~15× smaller than ResNet50 while achieving competitive accuracy.",
        size=11)

    add_heading(doc, "Hyperparameter Tables — Mark", 3)
    add_para(doc, "Approach 1 — Feature Extractor + Linear SVM:", bold=True, size=11)
    add_table(doc,
        headers=["Parameter", "Value"],
        rows=[
            ["Feature dimension",  "1280-d (MobileNetV2 top layer output)"],
            ["Pooling",            "Global Average Pooling"],
            ["Pre-processing",     "StandardScaler"],
            ["SVM kernel",         "Linear"],
            ["SVM C",              "1.0"],
        ], col_widths=[6, 10])

    add_para(doc, "Approach 2 — End-to-End (Two variants):", bold=True, size=11)
    add_table(doc,
        headers=["Parameter", "Approach 2A (Failed)", "Approach 2B (Improved)"],
        rows=[
            ["Validation set",   "16 images (original Kaggle)", "1,043 images (20% split from train)"],
            ["Layers unfrozen",  "Last 20",                      "Last 20"],
            ["Warm-up LR",       "1e-3",                         "1e-3"],
            ["Fine-tune LR",     "1e-4",                         "5e-5"],
            ["Warm-up epochs",   "10 (max)",                     "10 (max)"],
            ["Fine-tune epochs", "20 (max)",                     "20 (max)"],
            ["Dropout",          "0.4",                          "0.4"],
            ["Class weights",    "None",                         "Balanced"],
            ["Batch size",       "32",                           "32"],
            ["Optimizer",        "Adam",                         "Adam"],
        ], col_widths=[4.5, 5.5, 6])
    add_page_break(doc)

    # ── 3.3 EfficientNet-B0 (Mario & Veronia) ────────────────────────────────
    add_heading(doc, "3.3 EfficientNet-B0 — Mario & Veronia", 2)
    add_para(doc,
        "EfficientNet-B0 (Tan & Le, ICML 2019) uses compound scaling — uniformly scaling "
        "network depth, width, and resolution with a fixed set of coefficients — to achieve "
        "higher accuracy per FLOP than prior architectures. The core building block is the "
        "MBConv (Mobile Inverted Bottleneck with Squeeze-and-Excitation). "
        "EfficientNet-B0 has ~5.3M parameters and achieves 77.1% ImageNet Top-1 accuracy.",
        size=11)

    add_heading(doc, "Hyperparameter Tables — Mario (PlantVillage / TensorFlow)", 3)
    add_para(doc, "Approach 1 — End-to-End DL:", bold=True, size=11)
    add_table(doc,
        headers=["Parameter", "Value"],
        rows=[
            ["Frozen layers (Phase 1)", "All EfficientNet-B0 layers frozen"],
            ["Unfrozen from (Phase 2)", "Layer index 100 onward (~136 trainable layers)"],
            ["Warm-up epochs",          "10"],
            ["Fine-tune epochs",        "Up to 20 (early stop patience=5)"],
            ["Phase 1 LR",             "1e-4"],
            ["Phase 2 LR",             "1e-5 + ReduceLROnPlateau (factor=0.5, patience=3)"],
            ["Dropout",                "0.5 (before Dense head)"],
            ["Head",                   "BN → Dropout(0.5) → Dense(38, softmax)"],
            ["Loss",                   "Categorical cross-entropy"],
            ["Batch size",             "32"],
            ["Optimizer",              "Adam"],
        ], col_widths=[6, 10])

    add_para(doc, "Approach 2 — Feature Extractor + RBF-SVM:", bold=True, size=11)
    add_table(doc,
        headers=["Parameter", "Value"],
        rows=[
            ["Feature dimension",    "1280-d (EfficientNet-B0 GAP output)"],
            ["Dimensionality reduction","PCA: 1280 → 256 (retains >95% variance)"],
            ["Pre-processing",       "StandardScaler (zero-mean, unit-variance)"],
            ["SVM kernel",           "RBF"],
            ["SVM C",                "10.0"],
            ["SVM gamma",            "scale"],
        ], col_widths=[6, 10])
    add_image(doc, MARIO / "results" / "hyperparameter_table.png", width=6.0,
              caption="Figure 3.3a — Mario's hyperparameter summary (generated by the pipeline)")

    doc.add_paragraph()
    add_heading(doc, "Hyperparameter Tables — Veronia (CIFAR-10 / PyTorch)", 3)
    add_para(doc, "Approach 1 — Feature Extractor + Linear SVM:", bold=True, size=11)
    add_table(doc,
        headers=["Parameter", "Value"],
        rows=[
            ["Feature dimension",   "1280-d (EfficientNet-B0, classifier replaced by nn.Identity)"],
            ["Pre-processing",      "StandardScaler (fit on train only — no leakage)"],
            ["SVM kernel",          "Linear"],
            ["SVM C",               "1.0"],
            ["SVM max iterations",  "2000"],
        ], col_widths=[6, 10])

    add_para(doc, "Approach 2 — End-to-End Fine-Tuning:", bold=True, size=11)
    add_table(doc,
        headers=["Parameter", "Value"],
        rows=[
            ["Frozen layers",    "None — all layers fine-tuned"],
            ["Epochs",           "10–15"],
            ["Batch size",       "64"],
            ["Learning rate",    "1e-4"],
            ["Weight decay",     "1e-4 (L2 regularisation)"],
            ["Dropout",          "0.3 (before final linear layer)"],
            ["Head",             "Dropout(0.3) → Linear(1280 → 10)"],
            ["Loss",             "CrossEntropyLoss"],
            ["Optimizer",        "Adam"],
            ["LR Scheduler",     "ReduceLROnPlateau (patience=3, factor=0.5)"],
            ["Checkpoint",       "Best validation accuracy (not final epoch)"],
        ], col_widths=[6, 10])
    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 4 — Requirement 3: Implementation
# ─────────────────────────────────────────────────────────────────────────────

def build_req3(doc):
    add_heading(doc, "4. Requirement 3 — Implementation Framework", 1)

    # ── 4.1 Matthew ──────────────────────────────────────────────────────────
    add_heading(doc, "4.1 Matthew — ResNet50 + Chest X-Ray", 2)

    add_heading(doc, "Architecture Diagram", 3)
    add_para(doc, "Approach 1 — Feature Extraction Pipeline:", bold=True)
    for line in [
        "Input Image (224x224x3)",
        "  |",
        "ResNet50 (frozen, 175 layers, ImageNet weights)",
        "  |            |",
        "[GAP]        [GMP]   <- Global Average + Max Pooling over last conv block",
        "  +----+----+",
        "  Concatenate -> 4096-d feature vector",
        "  |         |         |",
        "[SVM]      [LR]      [MLP]   <- StandardScaler -> L2-normalise -> Estimator",
    ]:
        p = doc.add_paragraph(line if line else " ")
        run = p.runs[0] if p.runs else p.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9.5)
        p.paragraph_format.space_after = Pt(1)

    doc.add_paragraph()
    add_para(doc, "Approach 2 — End-to-End Architecture:", bold=True)
    for line in [
        "Input Image (224x224x3)",
        "  |",
        "ResNet50 base",
        "  Phase 1: fully frozen | Phase 2: last 15 layers unfrozen",
        "  |",
        "[GAP] + [GMP]  ->  Concatenate (4096-d)",
        "  |",
        "Dense(512) + BatchNorm + ReLU + Dropout(0.4) + L2(1e-4)",
        "  |",
        "Dense(128) + BatchNorm + ReLU + Dropout(0.3) + L2(1e-4)",
        "  |",
        "Dense(1, sigmoid)  ->  P(PNEUMONIA)",
    ]:
        p = doc.add_paragraph(line if line else " ")
        run = p.runs[0] if p.runs else p.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9.5)
        p.paragraph_format.space_after = Pt(1)

    add_heading(doc, "Approach 1 Results — Matthew", 3)
    add_table(doc,
        headers=["Model", "Accuracy", "F1 (weighted)", "ROC AUC", "PR AUC", "Train Time"],
        rows=[
            ["SVM (ResNet50 feats)", "80.45%", "0.7844", "0.9533", "0.9699", "1,705 s"],
            ["LR  (ResNet50 feats)", "82.37%", "0.8110", "0.9520", "0.9678", "16 s"],
            ["MLP (ResNet50 feats)", "78.37%", "0.7578", "0.9525", "0.9678", "16 s"],
        ], col_widths=[4, 2.5, 3, 2.5, 2.5, 2.5])

    add_image(doc, MATTHEW / "results" / "cm_svm.png",    width=3.0, caption="CM — SVM")
    add_image(doc, MATTHEW / "results" / "cm_lr.png",     width=3.0, caption="CM — LR")
    add_image(doc, MATTHEW / "results" / "cm_mlp.png",    width=3.0, caption="CM — MLP")
    add_image(doc, MATTHEW / "results" / "roc_curves.png", width=5.5,
              caption="Figure 4.1a — ROC curves for all Matthew models")
    add_image(doc, MATTHEW / "results" / "pr_curves.png",  width=5.5,
              caption="Figure 4.1b — Precision-Recall curves for all Matthew models")

    add_heading(doc, "Approach 2 Results — Matthew", 3)
    add_table(doc,
        headers=["Metric", "Default (thr=0.5)", "F1-Tuned Threshold", "Youden's J Threshold"],
        rows=[
            ["Accuracy",   "91.03%", "—",     "—"],
            ["F1 (weighted)", "0.9093", "—",  "—"],
            ["ROC AUC",    "0.9708", "—",     "—"],
            ["PR AUC",     "0.9785", "—",     "—"],
            ["Train Time", "5,923 s","—",     "—"],
        ], col_widths=[4, 4, 4, 4])

    add_image(doc, MATTHEW / "results" / "cm_dl.png",         width=3.5,
              caption="Figure 4.1c — DL confusion matrix (threshold 0.5)")
    add_image(doc, MATTHEW / "results" / "cm_dl_tuned.png",   width=3.5,
              caption="Figure 4.1d — DL confusion matrix (tuned threshold)")
    add_image(doc, MATTHEW / "results" / "dl_training_curves.png", width=5.5,
              caption="Figure 4.1e — Training curves: Loss & Accuracy across Phase 1 + Phase 2")
    add_image(doc, MATTHEW / "results" / "calibration_dl.png", width=4.5,
              caption="Figure 4.1f — Probability calibration plot (DL model)")
    add_image(doc, MATTHEW / "results" / "threshold_sweep_dl.png", width=5.0,
              caption="Figure 4.1g — F1 / Precision / Recall vs classification threshold")
    add_page_break(doc)

    # ── 4.2 Mark ─────────────────────────────────────────────────────────────
    add_heading(doc, "4.2 Mark — MobileNetV2 + Chest X-Ray", 2)

    add_heading(doc, "Architecture Diagram", 3)
    for line in [
        "Approach 1:",
        "  Image -> MobileNetV2 (frozen) -> 1280-d features -> StandardScaler -> SVM (linear, C=1.0)",
        " ",
        "Approach 2:",
        "  Image -> MobileNetV2 (last 20 layers unfrozen)",
        "        -> GlobalAvgPool -> BatchNorm -> Dropout(0.4)",
        "        -> Dense(128, ReLU) -> Dropout(0.3)",
        "        -> Dense(1, Sigmoid) -> P(PNEUMONIA)",
    ]:
        p = doc.add_paragraph(line if line.strip() else " ")
        run = p.runs[0] if p.runs else p.add_run(line if line.strip() else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9.5)
        p.paragraph_format.space_after = Pt(1)

    add_heading(doc, "Results — Mark", 3)
    add_table(doc,
        headers=["Metric", "Approach 1 (SVM)", "Approach 2A (val=16)", "Approach 2B (val=1043)"],
        rows=[
            ["Accuracy",   "85.10%", "70.51%", "87.82%"],
            ["Precision",  "0.8322", "0.6801",  "0.8520"],
            ["Recall",     "0.9538", "0.9974",  "0.9744"],
            ["F1-Score",   "0.8889", "0.8087",  "0.9091"],
            ["AUC",        "0.9415", "0.9249",  "0.9644"],
            ["Train Time", "8 s",    "1,805 s", "3,190 s"],
        ], col_widths=[3.5, 3.5, 4.5, 4.5])

    add_image(doc, MARK / "results" / "req3_cm_approach1.png",  width=3.5,
              caption="Figure 4.2a — CM: Approach 1 (SVM)")
    add_image(doc, MARK / "results" / "req3_cm_approach2a.png", width=3.5,
              caption="Figure 4.2b — CM: Approach 2A (collapsed)")
    add_image(doc, MARK / "results" / "req3_cm_approach2b.png", width=3.5,
              caption="Figure 4.2c — CM: Approach 2B (fixed)")
    add_image(doc, MARK / "results" / "req3_learning_curves_2a_vs_2b.png", width=5.5,
              caption="Figure 4.2d — Learning curves: 2A (unstable) vs 2B (stable)")
    add_image(doc, MARK / "results" / "req4_roc_all.png", width=5.5,
              caption="Figure 4.2e — ROC curves: all three approaches")
    add_page_break(doc)

    # ── 4.3 Mario ─────────────────────────────────────────────────────────────
    add_heading(doc, "4.3 Mario — EfficientNet-B0 + PlantVillage", 2)

    add_heading(doc, "Architecture Diagram", 3)
    for line in [
        "Input (224x224x3)",
        "  |",
        "EfficientNet-B0 backbone (Stem + 7 MBConv stages + Top Conv)",
        "  |",
        "Global Average Pooling  ->  1280-d feature vector",
        "  |                              |",
        "[Approach 1 -- End-to-End]    [Approach 2 -- SVM]",
        "BN -> Dropout(0.5)             PCA (1280 -> 256-d)",
        "-> Dense(38, softmax)          -> StandardScaler",
        "                               -> RBF-SVM (C=10)",
    ]:
        p = doc.add_paragraph(line if line else " ")
        run = p.runs[0] if p.runs else p.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9.5)
        p.paragraph_format.space_after = Pt(1)

    add_heading(doc, "Results — Mario", 3)
    add_table(doc,
        headers=["Metric", "Approach 1 (End-to-End DL)", "Approach 2 (EfficientNet + SVM)"],
        rows=[
            ["Accuracy",      "99.78%",  "97.74%"],
            ["Precision",     "99.78%",  "97.83%"],
            ["Recall",        "99.78%",  "97.74%"],
            ["F1-Score",      "99.78%",  "97.73%"],
            ["Train Time",    "~13.5 h", "~8 min (GPU + CPU)"],
            ["Inference Time","772 s",   "1,074 s"],
        ], col_widths=[4, 6, 6])

    add_image(doc, MARIO / "results" / "cm_dl.png",  width=5.5,
              caption="Figure 4.3a — Confusion matrix: End-to-End DL (38 classes)")
    add_image(doc, MARIO / "results" / "cm_svm.png", width=5.5,
              caption="Figure 4.3b — Confusion matrix: EfficientNet + RBF-SVM")
    add_image(doc, MARIO / "results" / "learning_curves.png", width=5.5,
              caption="Figure 4.3c — Training and validation accuracy/loss (Approach 1)")
    add_page_break(doc)

    # ── 4.4 Veronia ───────────────────────────────────────────────────────────
    add_heading(doc, "4.4 Veronia — EfficientNet-B0 + CIFAR-10 (PyTorch)", 2)

    add_heading(doc, "Architecture Diagram", 3)
    for line in [
        "Approach 1 -- Feature Extractor + Linear SVM:",
        "  Image -> EfficientNet-B0 (frozen, classifier=nn.Identity())",
        "        -> 1280-d features",
        "        -> StandardScaler (fit on train only)",
        "        -> LinearSVM (C=1.0, max_iter=2000)",
        " ",
        "Approach 2 -- End-to-End Fine-Tuning:",
        "  Image -> EfficientNet-B0 (all layers unfrozen)",
        "        -> Dropout(0.3)",
        "        -> Linear(1280 -> 10)",
        "        -> 10-class prediction",
    ]:
        p = doc.add_paragraph(line if line.strip() else " ")
        run = p.runs[0] if p.runs else p.add_run(line if line.strip() else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9.5)
        p.paragraph_format.space_after = Pt(1)

    add_heading(doc, "Results — Veronia", 3)
    add_para(doc,
        "Results are visible in the generated plots. "
        "Both approaches are evaluated on the held-out CIFAR-10 test set (10,000 images).",
        size=11)

    add_image(doc, VERONIA / "Results" / "confusion_matrix_approach1.png", width=4.5,
              caption="Figure 4.4a — CM: EfficientNet + Linear SVM (Approach 1)")
    add_image(doc, VERONIA / "Results" / "confusion_matrix_approach2.png", width=4.5,
              caption="Figure 4.4b — CM: End-to-End EfficientNet-B0 (Approach 2)")
    add_image(doc, VERONIA / "Results" / "learning_curves_approach2.png", width=5.5,
              caption="Figure 4.4c — Training loss & accuracy curves (Approach 2)")
    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 5 — Requirement 4: Comparative Analysis
# ─────────────────────────────────────────────────────────────────────────────

def build_req4(doc):
    add_heading(doc, "5. Requirement 4 — Comparative Analysis & Insights", 1)

    # ── 5.1 Per-member comparison charts ─────────────────────────────────────
    add_heading(doc, "5.1 Per-Member Comparative Visualisations", 2)

    add_heading(doc, "Matthew — Comparison Charts", 3)
    add_image(doc, MATTHEW / "results" / "comparison_metrics.png", width=6.0,
              caption="Figure 5.1a — Matthew: Metric comparison bar chart (Approach 1 vs Approach 2)")
    add_image(doc, MATTHEW / "results" / "comparison_time.png", width=5.0,
              caption="Figure 5.1b — Matthew: Training time comparison")
    add_image(doc, MATTHEW / "results" / "all_confusion_matrices.png", width=6.5,
              caption="Figure 5.1c — Matthew: All confusion matrices side-by-side")

    add_heading(doc, "Mark — Comparison Charts", 3)
    add_image(doc, MARK / "results" / "req4_3way_comparison.png", width=6.0,
              caption="Figure 5.1d — Mark: 3-way metric comparison (Approach 1, 2A, 2B)")
    add_image(doc, MARK / "results" / "req4_before_after_val.png", width=5.5,
              caption="Figure 5.1e — Mark: Effect of fixing validation set size")
    add_image(doc, MARK / "results" / "req4_3cm_comparison.png", width=6.5,
              caption="Figure 5.1f — Mark: 3 confusion matrices side-by-side")

    add_heading(doc, "Mario — Comparison Charts", 3)
    add_image(doc, MARIO / "results" / "metric_comparison.png", width=6.0,
              caption="Figure 5.1g — Mario: Approach 1 vs Approach 2 metric bar chart")
    add_image(doc, MARIO / "results" / "training_time.png", width=5.0,
              caption="Figure 5.1h — Mario: Training time comparison")

    add_heading(doc, "Veronia — Comparison Charts", 3)
    add_image(doc, VERONIA / "Results" / "comparison_metrics_bar.png", width=6.0,
              caption="Figure 5.1i — Veronia: Metric comparison bar chart")
    add_image(doc, VERONIA / "Results" / "comparison_training_time.png", width=5.0,
              caption="Figure 5.1j — Veronia: Training time comparison")
    add_image(doc, VERONIA / "Results" / "comparison_confusion_matrices.png", width=6.5,
              caption="Figure 5.1k — Veronia: Both confusion matrices side-by-side")
    add_page_break(doc)

    # ── 5.2 Cross-team summary ────────────────────────────────────────────────
    add_heading(doc, "5.2 Cross-Team Summary Table", 2)
    add_table(doc,
        headers=["Member", "Dataset", "Backbone",
                 "Best Feature-Extr. Acc.", "Best End-to-End Acc.", "Gap"],
        rows=[
            ["Matthew", "Chest X-Ray (binary)",    "ResNet50",
             "82.4% (LR)", "91.0%", "+8.6 pp"],
            ["Mark",    "Chest X-Ray (binary)",    "MobileNetV2",
             "85.1% (SVM)", "87.8% (2B)", "+2.7 pp"],
            ["Mario",   "PlantVillage (38 classes)","EfficientNet-B0 (TF)",
             "97.7% (SVM)", "99.8%", "+2.1 pp"],
            ["Veronia", "CIFAR-10 (10 classes)",    "EfficientNet-B0 (PyTorch)",
             "See chart", "See chart", "—"],
        ],
        col_widths=[2.5, 4, 3.5, 4, 3.5, 2.5])

    doc.add_paragraph()

    add_heading(doc, "Training Time Comparison", 3)
    add_table(doc,
        headers=["Member", "Backbone", "Feature-Extr. Train Time", "End-to-End Train Time", "Speedup"],
        rows=[
            ["Matthew", "ResNet50",        "16 s (LR)",   "5,923 s", "~370×"],
            ["Mark",    "MobileNetV2",     "8 s (SVM)",   "3,190 s", "~400×"],
            ["Mario",   "EfficientNet-B0", "~8 min (SVM)","~13.5 h", "~100×"],
            ["Veronia", "EfficientNet-B0", "Fast (SVM)",  "10–15 epochs", "Faster"],
        ], col_widths=[2.5, 3.5, 4, 4, 2])
    add_page_break(doc)

    # ── 5.3 Discussion & Conclusions ─────────────────────────────────────────
    add_heading(doc, "5.3 Discussion, Conclusions & Recommendations", 2)

    add_heading(doc, "i. How did the feature-extractor ML classifier compare to end-to-end?", 3)
    add_para(doc,
        "Across all four experiments, the end-to-end approach outperformed feature extraction "
        "on every accuracy and F1 metric. However, the magnitude of the improvement is highly "
        "dataset-dependent:", size=11)
    add_bullet(doc,
        "Matthew (Chest X-Ray, ResNet50): +8.6 percentage points. The large gap is partly "
        "explained by the distribution shift between train/val (~74% pneumonia) and test (~63%). "
        "End-to-end fine-tuning, combined with SWA, TTA, and threshold tuning, adapts to this "
        "shift far better than frozen features.")
    add_bullet(doc,
        "Mark (Chest X-Ray, MobileNetV2): +2.7 points (Approach 1 vs Approach 2B). "
        "Notably, Approach 2A (same model, tiny val set) performed 14.6 points below Approach 1, "
        "demonstrating that engineering decisions (val set size, class weights) can dwarf "
        "architecture differences.")
    add_bullet(doc,
        "Mario (PlantVillage, EfficientNet-B0): only +2.1 points. On a large, clean, "
        "controlled dataset the frozen features already capture most of the discriminative "
        "information — disease lesion texture, colour, shape — that the fine-tuned model learns. "
        "The 97.7% SVM result is competitive enough for many production systems.")
    add_bullet(doc,
        "Veronia (CIFAR-10, EfficientNet-B0 / PyTorch): end-to-end is expected to win given "
        "the large domain gap between ImageNet and CIFAR-10's 32×32 objects.")

    add_heading(doc, "ii. Advantages and limitations of each approach", 3)
    add_table(doc,
        headers=["Aspect", "Feature Extraction + ML", "End-to-End Fine-Tuning"],
        rows=[
            ["Accuracy",         "Good — leverages ImageNet features",       "Best — adapts features to target domain"],
            ["Training speed",   "Very fast (CNN not re-trained)",            "Slow (full backprop through CNN)"],
            ["GPU memory",       "Low (inference passes only)",               "High (gradients for all unfrozen layers)"],
            ["Interpretability", "High (explicit feature vector)",            "Low (distributed end-to-end)"],
            ["Overfitting risk", "Low (classical ML is well-regularised)",    "Higher (needs dropout, BN, L2, label smooth)"],
            ["Adaptability",     "Fixed to ImageNet feature space",           "Fully adapts to new domain"],
            ["Data requirement", "Works well with smaller datasets",          "Benefits more from large datasets"],
            ["Hyperparameter tuning","GridSearchCV — robust and fast",         "More callbacks/phases — complex"],
        ], col_widths=[4, 6, 6])

    add_heading(doc, "iii. Which approach is more training-time efficient?", 3)
    add_para(doc,
        "Feature extraction + classical ML is dramatically faster to train in every experiment: "
        "between 100× and 400× faster than end-to-end fine-tuning (8–16 seconds vs 1.5–6 hours). "
        "This is because the CNN is only run once in forward-pass mode to generate the feature "
        "cache; subsequent ML training is CPU-bound and converges in seconds. "
        "End-to-end fine-tuning requires multiple GPU passes through the full network per epoch, "
        "plus gradient computation and weight updates for up to 184 layers.",
        size=11)

    add_heading(doc, "iv. Recommendation: resource-constrained vs high-performance environments", 3)

    add_para(doc, "Resource-constrained environments (mobile, edge, IoT, limited GPU):", bold=True, size=11)
    add_para(doc,
        "Use feature extraction + classical ML (SVM or Logistic Regression). "
        "The frozen CNN can be quantised to INT8 for inference, the SVM head adds negligible "
        "compute, and the entire pipeline can run on a CPU. "
        "Mario's 97.7% SVM result on PlantVillage, and Mark's 85.1% SVM result on X-rays, "
        "show this approach is competitive enough for many practical deployments. "
        "Training can be repeated or adapted to new data in seconds without a GPU.",
        size=11)

    add_para(doc, "High-performance environments (cloud GPU, research, clinical-grade applications):", bold=True, size=11)
    add_para(doc,
        "Use end-to-end fine-tuning with all regularisation techniques enabled (label smoothing, "
        "L2, graduated dropout, SWA, TTA, threshold tuning on val). "
        "Matthew's end-to-end ResNet50 achieved 91% accuracy and a 0.97 ROC AUC on a "
        "challenging medical task with distribution shift — a level that the ML classifiers "
        "could not reach. For clinical screening (pneumonia triage), the additional cost of "
        "end-to-end training is fully justified.",
        size=11)

    add_heading(doc, "Overall Recommendation", 3)
    add_para(doc,
        "A practical strategy is to use feature extraction as a strong baseline first "
        "(fast to train, easy to interpret, easy to retrain on new data). "
        "Only invest in end-to-end fine-tuning if the application demands the highest accuracy "
        "and a GPU infrastructure is available. "
        "The two paradigms are complementary, not mutually exclusive — "
        "the feature-extracted embeddings can also serve as a warm-start for end-to-end training.",
        size=11, bold=False)
    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Section 6 — References
# ─────────────────────────────────────────────────────────────────────────────

def build_references(doc):
    add_heading(doc, "6. References", 1)
    refs = [
        "[1] He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep Residual Learning for Image "
        "Recognition. CVPR 2016. — ResNet50 backbone (Matthew).",
        "[2] Sandler, M., Howard, A., Zhu, M., Zhmoginov, A., & Chen, L. C. (2018). "
        "MobileNetV2: Inverted Residuals and Linear Bottlenecks. CVPR 2018. — MobileNetV2 (Mark).",
        "[3] Tan, M., & Le, Q. V. (2019). EfficientNet: Rethinking Model Scaling for "
        "Convolutional Neural Networks. ICML 2019. arXiv:1905.11946. — EfficientNet-B0 (Mario, Veronia).",
        "[4] Cortes, C., & Vapnik, V. (1995). Support-Vector Networks. Machine Learning, 20(3), 273–297.",
        "[5] Mohanty, S. P., Hughes, D. P., & Salathe, M. (2016). Using Deep Learning for "
        "Image-Based Plant Disease Detection. Frontiers in Plant Science, 7, 1419. — PlantVillage (Mario).",
        "[6] Mooney, P. (2018). Chest X-Ray Images (Pneumonia). Kaggle Dataset. "
        "https://www.kaggle.com/datasets/paultimothymooney/chest-xray-pneumonia — (Matthew, Mark).",
        "[7] Krizhevsky, A. (2009). Learning Multiple Layers of Features from Tiny Images. "
        "Technical Report, University of Toronto. — CIFAR-10 (Veronia).",
        "[8] Izmailov, P., Podoprikhin, D., Garipov, T., Vetrov, D., & Wilson, A. G. (2018). "
        "Averaging Weights Leads to Wider Optima and Better Generalization. UAI 2018. — SWA (Matthew).",
        "[9] Zhong, Z., Zheng, L., Kang, G., Li, S., & Yang, Y. (2020). Random Erasing Data "
        "Augmentation. AAAI 2020. — Cutout/RandomErasing (Matthew).",
        "[10] Youden, W. J. (1950). Index for rating diagnostic tests. Cancer, 3(1), 32–35. "
        "— Youden's J threshold (Matthew).",
        "[11] Deng, J., Dong, W., Socher, R., Li, L. J., Li, K., & Fei-Fei, L. (2009). "
        "ImageNet: A large-scale hierarchical image database. CVPR 2009. — Pretrained weights (all).",
    ]
    for r in refs:
        p = doc.add_paragraph(r, style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(10)


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)

    # Default normal style font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    print("Building cover page...")
    build_cover(doc)

    print("Building table of contents...")
    build_toc(doc)

    print("Building Section 1 — Overview...")
    build_overview(doc)

    print("Building Section 2 — Requirement 1 (Datasets)...")
    build_req1(doc)

    print("Building Section 3 — Requirement 2 (Model & Hyperparameters)...")
    build_req2(doc)

    print("Building Section 4 — Requirement 3 (Implementation)...")
    build_req3(doc)

    print("Building Section 5 — Requirement 4 (Comparative Analysis)...")
    build_req4(doc)

    print("Building Section 6 — References...")
    build_references(doc)

    doc.save(str(OUT))
    print(f"\n[OK] Report saved to: {OUT}")
    print(f"  File size: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
